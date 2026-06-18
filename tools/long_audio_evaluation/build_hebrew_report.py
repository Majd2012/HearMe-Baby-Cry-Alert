from __future__ import annotations

import json
import os
import sys
from pathlib import Path

os.environ.setdefault("MPLCONFIGDIR", str(Path.cwd() / ".matplotlib"))

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from detector_core import read_ground_truth
from scientific_validation import (
    CURRENT_SELECTED_CONFIG,
    ORIGINAL_CONFIG,
    create_zoom_plots,
    load_config,
    apply_policy,
    evaluate_events,
    row_to_policy,
)
from generate_state_chart import main as generate_state_charts


ROOT = Path(__file__).resolve().parents[2]
VALIDATION_DIR = ROOT / "results" / "scientific_validation"
PLOTS_DIR = VALIDATION_DIR / "plots"
OUT_DIR = ROOT / "results" / "final_report"
DOCX_PATH = OUT_DIR / "final_hebrew_algorithm_report.docx"


TITLE = "מערכת לזיהוי בכי תינוק והתרעה לשעון חכם – ניתוח אלגוריתמי ובחירת נקודת עבודה מיטבית"


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    regenerate_true_noncry_zoom()
    generate_state_charts()
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    add_intro(doc)
    add_initial_algorithm(doc)
    add_professor_concern(doc)
    add_long_audio_system(doc)
    add_synthetic_night_section(doc)
    add_alert_logic(doc)
    add_config_sections(doc)
    add_results_sections(doc)
    add_figures(doc)
    add_limitations_and_conclusion(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


def setup_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    for style_name in ["Normal", "Body Text"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(11)
    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def rtl_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)


def add_p(doc, text="", style=None, bold=False):
    p = doc.add_paragraph(style=style)
    rtl_paragraph(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.bold = bold
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    rtl_paragraph(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        rtl_paragraph(p)
        run = p.add_run(item)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def add_cover(doc):
    p = doc.add_paragraph()
    rtl_paragraph(p)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(TITLE)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(13, 40, 77)
    add_p(doc, "דוח מסכם המיועד להצגה אקדמית: הערכת האלגוריתם, בדיקה על שמע רציף, והשוואה לפני ואחרי אופטימיזציה.", bold=True)
    add_p(doc, "המערכת: HearMe – זיהוי בכי תינוק בטלפון ושליחת התרעת רטט לשעון חכם.")
    add_p(doc, "מקורות הנתונים בדוח זה: קבצי התוצאה שנוצרו בתיקיית results/scientific_validation.")
    doc.add_page_break()


def add_intro(doc):
    add_heading(doc, "1. מבוא", 1)
    add_p(doc, "מטרת הפרויקט היא לסייע להורים כבדי שמיעה לזהות בזמן אמת מצב שבו התינוק בוכה. הטלפון הנייד מונח ליד התינוק, מאזין לסביבה, מנתח את האודיו, וכאשר מזוהה אירוע בכי אמין הוא שולח התרעה לשעון חכם. השעון רוטט וכך מודיע להורה שיש צורך לבדוק את התינוק.")
    add_p(doc, "המערכת מחולקת לשני חלקים עיקריים: אפליקציית Android בטלפון שאחראית על הקלטת האודיו, הרצת המודל וקבלת ההחלטה, ואפליקציית Wear OS בשעון שמקבלת את ההודעה ומפעילה רטט והתראה.")


def add_initial_algorithm(doc):
    add_heading(doc, "2. האלגוריתם הראשוני", 1)
    add_p(doc, "בשלב הראשון השתמשנו ב-YAMNet, מודל סיווג אודיו מאומן מראש של Google. המודל מספק ציון הסתברותי למחלקות קול שונות, ובפרויקט נבחר הציון של המחלקה Baby Cry / Infant Cry. בסביבת המחקר המודל הופעל דרך TensorFlow, ובאפליקציה דרך TensorFlow Lite.")
    add_p(doc, "בתחילת העבודה נבדקו קטעי אודיו נפרדים ומתוייגים. עבור כל קטע חושב ציון בכי, ולאחר מכן סף החלטה קבע האם הקטע מסווג כ-Cry או No Cry.")
    add_bullets(doc, [
        "מספר דגימות כולל: 2,353",
        "Accuracy: 85.5%",
        "Cry precision: 98.7%",
        "Cry recall: 84.3%",
    ])
    add_p(doc, "תוצאות אלו הראו שהמודל מתאים כבסיס, אך הן אינן מספיקות כדי להוכיח ביצועים באפליקציה אמיתית, מכיוון שהאפליקציה עובדת על שמע רציף ולא על קבצים מופרדים מראש.")


def add_professor_concern(doc):
    add_heading(doc, "3. הערת המרצה והצורך בהערכה רציפה", 1)
    add_p(doc, "המרצה הדגיש כי תוצאות על קטעי אודיו נפרדים אינן בהכרח מייצגות את ביצועי האפליקציה בזמן אמת. באפליקציה קיימים שלבים נוספים שמשפיעים על ההחלטה: חלוקה לחלונות זמן, ספי החלטה, החלקת ציונים, דרישת התמדה, מניעת התרעות חוזרות, זמן קירור והפעלה מחדש של הגלאי.")
    add_p(doc, "לכן נבנתה מערכת הערכה חדשה עבור קובצי אודיו ארוכים. מערכת זו מאפשרת לבדוק מתי התינוק באמת התחיל לבכות, מתי המודל חצה את הסף, מתי נשלחה התרעה בפועל, והאם היו התרעות שווא או אירועים שהוחמצו.")


def add_long_audio_system(doc):
    add_heading(doc, "4. מערכת הערכה לשמע ארוך", 1)
    add_p(doc, "נוצרו הקלטות לילה סינתטיות המשלבות רעשי רקע של חדר עם אירועי בכי שהוכנסו במיקומים ידועים. לכל אירוע נשמרים זמן התחלה וזמן סיום בקובץ ground-truth CSV. כך ניתן להשוות בין המציאות הידועה לבין ההתרעות שיצרה המערכת.")
    add_bullets(doc, [
        "מספר לילות סינתטיים בתיקוף הסופי: 3",
        "משך כל לילה: 30 דקות",
        "משך בדיקה כולל: 1.5 שעות",
        "מספר אירועי בכי כולל: 18",
    ])
    add_p(doc, "המדדים שחושבו כוללים recall ברמת אירוע, precision ברמת אירוע, F1, מספר אירועים שהוחמצו, התרעות שווא לשעה, התרעות מוקדמות, התרעות כפולות וחציון זמן ההשהיה עד התרעה.")


def add_synthetic_night_section(doc):
    add_heading(doc, "5. יצירת הקלטת לילה סינתטית", 1)
    add_p(doc, "כדי לבדוק את האלגוריתם בצורה מבוקרת יצרנו הקלטת לילה סינתטית בעזרת סקריפט Python ייעודי. כלומר, לא הסתפקנו בבדיקת קטעי בכי קצרים ונפרדים, אלא בנינו קובץ אודיו ארוך המדמה מצב שבו טלפון נמצא ליד התינוק במשך זמן ממושך ומאזין לרעשי סביבה ולבכי.")
    add_p(doc, "הסיבה לשימוש בהקלטה סינתטית היא שלא היה בידינו קובץ לילה אמיתי מלא שבו זמני התחלה וסיום של כל אירוע בכי אומתו במדויק. בהקלטה סינתטית אנו שולטים בזמן שבו כל בכי מוכנס, ולכן אנו יודעים את האמת הקרקעית (ground truth) של כל אירוע.")
    add_p(doc, "הקלטה סינתטית אינה מיועדת לטעון שמדובר בלילה אמיתי מלא, אלא ליצור סביבת בדיקה מבוקרת שבה ידועה האמת הקרקעית (ground truth) של כל אירוע בכי.")
    add_p(doc, "קובצי הקלט שנעשה בהם שימוש היו משני סוגים: הקלטות רקע של חדר, וקבצי אודיו של בכי תינוק. רעשי הרקע שימשו כבסיס רציף המדמה סביבה ביתית, ודגימות הבכי הוכנסו לתוך הרקע בזמנים ידועים.")
    add_p(doc, "בפועל, הסקריפט קיבל כקלט את קובצי הרקע ואת תיקיית דגימות הבכי, ולאחר מכן יצר קובץ אודיו ארוך אחד. התהליך בוצע באופן אוטומטי כדי שכל ניסוי יהיה ניתן לשחזור.")
    add_heading(doc, "5.1 תהליך יצירת ההקלטה", 2)
    add_p(doc, "תהליך הבנייה של ההקלטה הסינתטית בוצע בשלבים הבאים:")
    add_bullets(doc, [
        "אודיו הרקע של החדר חזר על עצמו או שורשר עד שהגיע למשך ההקלטה המבוקש.",
        "דגימות בכי תינוק הוכנסו לתוך הרקע בזמני התחלה מבוקרים.",
        "זמני ההכנסה נוצרו באמצעות random seed, ולכן ניתן לשחזר את הניסוי.",
        "אירועי הבכי קיבלו משכים שונים ורמות עוצמה שונות כדי לדמות מקרים קלים ומקרים קשים לזיהוי.",
        "נוספו fade-in ו-fade-out קצרים כדי למנוע חיתוכי אודיו לא טבעיים.",
        "רמות האודיו הותאמו כדי למנוע clipping.",
    ])
    add_heading(doc, "5.2 קובצי הפלט שנוצרו", 2)
    add_p(doc, "פלט תהליך היצירה כלל שלושה קבצים מרכזיים:")
    add_bullets(doc, [
        "synthetic_night.wav – קובץ האודיו הארוך הסופי ששימש לבדיקה.",
        "synthetic_night_ground_truth.csv – קובץ המכיל את זמני ההתחלה והסיום האמיתיים של כל אירוע בכי שהוכנס.",
        "synthetic_night_manifest.json – קובץ תיעוד הכולל את קובצי המקור, ה-random seed, זמני ההכנסה, משכי האירועים והגדרות gain / loudness.",
    ])
    add_heading(doc, "5.3 מדוע השיטה שימושית מבחינה מדעית", 2)
    add_p(doc, "החשיבות המדעית של שיטה זו היא שהיא מאפשרת להשוות בין זמן הבכי האמיתי לבין זמן ההתרעה, לחשב detection latency, למדוד אירועים שהוחמצו, לזהות התרעות שווא ולבדוק האם קיימות התרעות כפולות. בנוסף, מכיוון שהניסוי מבוסס על seed ותיעוד מלא, ניתן לשחזר אותו.")
    add_p(doc, "אם מריצים את אותו הסקריפט עם אותו random seed ואותם קובצי מקור, מתקבלת אותה הקלטה ואותם זמני בכי. לכן הניסוי אינו תלוי בהחלטה ידנית או בעריכה חד-פעמית בתוכנת מיקס, אלא בתהליך מסודר שניתן להסביר ולחזור עליו.")
    add_heading(doc, "5.4 זרימת העבודה", 2)
    add_p(doc, "Room background audio + Baby cry samples → Synthetic night WAV → Ground-truth CSV → Long-audio evaluation → Metrics and plots")
    add_p(doc, "במילים פשוטות: יצרנו קובץ שמע ארוך עם רעשי רקע, הכנסנו לתוכו בכי תינוק בזמנים ידועים, שמרנו את הזמנים האמיתיים בקובץ CSV, ואז בדקנו האם האלגוריתם התריע בזמן הנכון.")


def add_alert_logic(doc):
    add_heading(doc, "6. שיפור לוגיקת ההתרעה", 1)
    add_p(doc, "כדי למנוע התרעה בעקבות קפיצה רגעית בציון המודל, הוגדרה מכונת מצבים מפורשת:")
    add_bullets(doc, ["IDLE", "POSSIBLE_CRY", "CONFIRMED_CRY", "ALERTED", "COOLDOWN", "REARMING"])
    add_p(doc, "מצב POSSIBLE_CRY מאפשר לאסוף ראיות לפני התרעה. לאחר אישור בכי המערכת נכנסת למצב ALERTED ולאחר מכן COOLDOWN, כדי שלא להעיר את ההורה שוב ושוב במהלך אותו אירוע בכי רציף. מצב REARMING דורש שהציון יחזור להיות נמוך לפני שהמערכת מוכנה לאירוע חדש.")
    chart_path = OUT_DIR / "state_chart_hebrew.png"
    if not chart_path.exists():
        raise FileNotFoundError(chart_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(chart_path), width=Inches(6.2))
    add_p(doc, "איור: מכונת המצבים של לוגיקת ההתרעה החדשה.", bold=True)
    add_p(doc, "התרשים מראה שהמערכת כבר אינה שולחת התרעה על סמך ציון גבוה יחיד. במקום זאת היא עוברת ברצף מבוקר של מצבים: חשד לבכי, אישור באמצעות התמדה, שליחת התרעה, זמן קירור והכנה מחדש. מבנה זה מפחית יקיצות שווא ומונע התרעות כפולות כאשר אותו אירוע בכי רציף עדיין נמשך.")


def add_config_sections(doc):
    add_heading(doc, "7. תצורת המערכת המקורית", 1)
    add_p(doc, "התצורה המקורית אומתה מהיסטוריית המאגר. הגרסה הראשונית ביותר השתמשה בסף 0.30 ובכלל של 20 מתוך 24 מקטעים. לצורך השוואת לפני/אחרי השתמשנו בתצורת מכונת המצבים שלפני האופטימיזציה, משום שהיא קודמת ישירה לתצורה הסופית:")
    add_bullets(doc, [
        "Trigger threshold: 0.30",
        "Clear threshold: 0.20",
        "Smoothing: rolling mean",
        "Persistence: 3 app segments",
        "Cooldown: 120 seconds",
        "Rearming: 5 seconds",
    ])
    add_heading(doc, "8. התצורה המומלצת הסופית", 1)
    add_p(doc, "התצורה שנבחרה לאחר תיקוף על שלושה לילות סינתטיים היא:")
    add_bullets(doc, [
        "Trigger threshold: 0.05",
        "Clear threshold: 0.03",
        "Persistence: 2 frames",
        "Smoothing: rolling mean, 3 frames",
        "Cooldown: 30 seconds",
        "Rearming: 5 seconds",
    ])
    add_p(doc, "התצורה שנבחרה קודם לכן הייתה רגישה מדי: סף 0.05, סף ניקוי 0.00, התמדה של פריים אחד וללא החלקה. התצורה הסופית יציבה יותר משום שהיא כוללת סף ניקוי מעל אפס, דרישת התמדה של שני פריימים, החלקת rolling mean, cooldown ו-rearming.")


def add_results_sections(doc):
    add_heading(doc, "9. תוצאות סופיות ועקרון האופטימיזציה", 1)
    add_p(doc, "עקרון הבחירה לא היה השגת 100% דיוק. מטרת הפרויקט היא למצוא נקודת עבודה מעשית המתאימה לסדר העדיפויות: קודם כל להימנע מהערת ההורים ללא סיבה אמיתית, לאחר מכן לצמצם החמצת בכי, לשמור על השהיית זיהוי סבירה, ולמנוע התרעות כפולות באותו אירוע.")
    add_p(doc, "ספים נמוכים יכולים לשפר recall אך עלולים ליצור התרעות שווא. ספים גבוהים מפחיתים התרעות שווא אך עלולים להחמיץ בכי חלש. לכן נבחרה נקודת עבודה שמאזנת בין הדרישות.")
    add_p(doc, "נקודת העבודה שנבחרה היא פשרה מעשית המתאימה לדרישות הפרויקט ולנתונים שנבדקו.")
    add_results_table(doc)
    add_p(doc, "אף תצורה שנבדקה לא הגיעה ל-recall של 85% או 90% תוך שמירה על אפס התרעות שווא. לכן נבחרה התצורה הקרובה ביותר שעמדה בעדיפות המרכזית: אפס התרעות שווא, עם recall הגבוה ביותר שנמצא תחת מגבלה זו.")


def add_results_table(doc):
    df = pd.read_csv(VALIDATION_DIR / "multi_night_parameter_sweep.csv")
    rows = []
    labels = {
        "original": "מערכת מקורית",
        "current_selected": "תצורה רגישה קודמת",
        "candidate_0024": "תצורה מומלצת סופית",
    }
    for key in ["original", "current_selected", "candidate_0024"]:
        row = df[df.config_name == key].iloc[0]
        rows.append([
            labels[key],
            f"{row.event_recall * 100:.1f}%",
            f"{row.event_precision * 100:.1f}%",
            f"{row.event_f1 * 100:.1f}%",
            f"{row.false_alerts_per_hour:.2f}",
            f"{int(row.missed_crying_events)}",
            f"{row.median_detection_latency:.2f}",
            f"{int(row.early_alerts)}",
            f"{int(row.duplicate_alerts)}",
        ])
    headers = ["תצורה", "Recall", "Precision", "F1", "False alerts/hour", "Missed cries", "Median latency (s)", "Early alerts", "Duplicate alerts"]
    add_table(doc, headers, rows)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        shade_cell(cell, "E8EEF5")
        set_cell_rtl(cell, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            set_cell_rtl(cells[i])
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_rtl(cell, bold=False):
    for paragraph in cell.paragraphs:
        rtl_paragraph(paragraph)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            run.bold = bold


def add_figures(doc):
    add_heading(doc, "10. תרשימים והסברים", 1)
    figures = [
        ("before_after_metrics.png", "השוואת מדדים לפני ואחרי האופטימיזציה", "בתרשים ניתן לראות שה-recall השתפר לאחר האופטימיזציה, מספר אירועי הבכי שהוחמצו ירד, חציון זמן הזיהוי ירד, ומספר התרעות השווא לשעה נשאר אפס. כלומר, המערכת השתפרה בזיהוי ועדיין שמרה על העדיפות המרכזית: לא להעיר את ההורים ללא סיבה."),
        ("best_operating_point.png", "בחירת נקודת העבודה המומלצת", "כל נקודה מייצגת תצורה שנבדקה. ציר ה-X מציג התרעות שווא לשעה וציר ה-Y מציג recall ברמת אירוע. הנקודה המומלצת נבחרה משום שהיא שומרת על אפס התרעות שווא ומשיגה את ה-recall הטוב ביותר שנמצא תחת מגבלה זו. זו פשרה מעשית ולא פתרון מתמטי מושלם."),
        ("timeline_before_modifications.png", "ציר זמן לפני שיפור לוגיקת ההתרעה", "האזורים הוורודים הם אירועי בכי אמיתיים, העקומות מציגות את ציון הבכי, קווי הסף מציגים את גבולות ההחלטה, הקווים הסגולים מציגים התרעות, וסימוני X מציינים אירועים שהוחמצו. לפני האופטימיזציה המערכת החמיצה אירועים רבים בגלל סף ולוגיקה שאינם מתאימים מספיק לשמע רציף."),
        ("timeline_after_optimization.png", "ציר זמן אחרי שיפור לוגיקת ההתרעה", "לאחר האופטימיזציה המערכת מזהה יותר אירועי בכי. הלוגיקה רגישה יותר אך עדיין מבוקרת בעזרת החלקה, התמדה ו-rearming. מספר התרעות השווא נשאר אפס, ועדיין חלק מאירועי הבכי החלשים הוחמצו."),
        ("zoom_correct_cry_before.png", "תקריב: אירוע בכי מזוהה לפני האופטימיזציה", "תקריב זה מציג אירוע בכי לפני האופטימיזציה. המערכת הצליחה לזהות את הבכי, אך ההתרעה הגיעה מאוחר יותר משום שהתצורה הקודמת דרשה ראיה חזקה ומתמשכת יותר."),
        ("zoom_correct_cry_after.png", "תקריב: אירוע בכי מזוהה אחרי האופטימיזציה", "בתצורה המשופרת המערכת מגיבה מוקדם יותר, אך עדיין משתמשת בהחלקה ובהתמדה כדי להימנע מתגובה לקפיצות אקראיות בציון."),
        ("zoom_difficult_or_missed_cry_before.png", "תקריב: אירוע קשה או מוחמץ לפני האופטימיזציה", "תרשים זה מציג אירוע בכי חלש או רועש שבו ציוני הבכי נמוכים יחסית. לפני האופטימיזציה הציון לא עבר את הסף המקורי ולכן האירוע הוחמץ."),
        ("zoom_difficult_or_missed_cry_after.png", "תקריב: אירוע קשה או מוחמץ אחרי האופטימיזציה", "גם לאחר האופטימיזציה קיימים אירועים חלשים שעלולים להיות מוחמצים. זהו חלק חשוב מהדוח, משום שהוא מראה שאיננו טוענים לביצועים מושלמים."),
        ("zoom_non_cry_section_before.png", "תקריב: קטע ללא בכי לפני האופטימיזציה", "קטע זה אומת כקטע ללא ground-truth של בכי. הוא מאפשר לבדוק האם המערכת יוצרת התרעת שווא בזמן שאין בכי אמיתי."),
        ("zoom_non_cry_section_after.png", "תקריב: קטע ללא בכי אחרי האופטימיזציה", "גם לאחר האופטימיזציה הקטע ללא בכי אינו אמור ליצור התרעה. הדבר חשוב במיוחד משום שהמרצה הדגיש שיש להימנע מהערת ההורים ללא סיבה אמיתית."),
    ]
    for filename, caption, explanation in figures:
        path = PLOTS_DIR / filename
        if not path.exists():
            raise FileNotFoundError(path)
        add_heading(doc, caption, 2)
        p = doc.add_paragraph()
        rtl_paragraph(p)
        r = p.add_run()
        r.add_picture(str(path), width=Inches(6.2))
        add_p(doc, explanation)


def add_limitations_and_conclusion(doc):
    add_heading(doc, "11. מגבלות", 1)
    add_bullets(doc, [
        "התיקוף מבוסס על הקלטות סינתטיות ולא על לילה אמיתי מלא.",
        "משך הבדיקה הכולל הוא 1.5 שעות בלבד.",
        "נבדקו 18 אירועי בכי.",
        "חלק מאירועי הבכי החלשים עדיין הוחמצו.",
        "יש להמשיך להשוות בין Python לבין האפליקציה על אותם קבצי אודיו.",
        "נדרשים יותר רעשי רקע והקלטות אמיתיות לפני טענה לאמינות ברמת מוצר.",
    ])
    add_heading(doc, "12. סיכום", 1)
    add_p(doc, "בפרויקט עברנו מהערכה על קבצי אודיו מופרדים להערכה רציפה על שמע ארוך. הוספנו מכונת מצבים ולוגיקת התרעה שמטפלת בהתמדה, cooldown ו-rearming. השווינו בין המערכת המקורית לבין המערכת המשופרת, ובחרנו נקודת עבודה שמעדיפה הימנעות מהתרעות שווא תוך שמירה על recall גבוה ככל האפשר תחת מגבלה זו.")
    add_p(doc, "המערכת הסופית אינה מושלמת, אך היא מציאותית יותר ומותאמת יותר לדרישות המרצה. נדרש תיקוף נוסף על הקלטות אמיתיות וארוכות לפני שימוש מעשי.")


def regenerate_true_noncry_zoom():
    gt = read_ground_truth(VALIDATION_DIR / "nights" / "synthetic_night_seed_42_ground_truth.csv", 1800.0)
    raw = pd.read_csv(VALIDATION_DIR / "frame_scores" / "seed_42_raw_frame_scores.csv")
    final_df = pd.read_csv(VALIDATION_DIR / "multi_night_parameter_sweep.csv")
    final_config = row_to_policy(final_df[final_df.config_name == "candidate_0024"].iloc[0])
    before, _ = apply_policy(raw, ORIGINAL_CONFIG)
    after, _ = apply_policy(raw, final_config)
    before.attrs["trigger_threshold"] = ORIGINAL_CONFIG.trigger_threshold
    after.attrs["trigger_threshold"] = final_config.trigger_threshold
    before_events, _ = evaluate_events(before, gt, 1800.0, 0.0, 5.0, "original")
    after_events, _ = evaluate_events(after, gt, 1800.0, 0.0, 5.0, "optimized")
    y_limits = (0.0, 0.4)
    create_zoom_plots(before, after, gt, before_events, after_events, ORIGINAL_CONFIG, final_config, PLOTS_DIR, y_limits)


if __name__ == "__main__":
    main()
